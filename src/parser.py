import os
import io
import logging
import docx

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Fallback imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF (fitz) is not available.")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber is not available.")

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("pypdf is not available.")

try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    logger.warning("pytesseract or PIL is not available. OCR fallback disabled.")

# Constrain all document ingestion to a strict 3-page processing limit
MAX_PAGES_TO_PARSE = 3

def extract_text_via_ocr(file_bytes):
    """
    Renders PDF pages into images using PyMuPDF and runs Tesseract OCR.
    Constrained to first 3 pages.
    """
    if not (PYMUPDF_AVAILABLE and PYTESSERACT_AVAILABLE):
        raise RuntimeError("PyMuPDF or pytesseract/Pillow is not available. OCR cannot be performed.")
    
    logger.info(f"Attempting local OCR via PyMuPDF page rendering + pytesseract (capped at first {MAX_PAGES_TO_PARSE} pages)...")
    text = ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_to_parse = min(MAX_PAGES_TO_PARSE, len(doc))
        
        for page_num in range(pages_to_parse):
            logger.info(f"OCRing page {page_num + 1}/{pages_to_parse}...")
            page = doc[page_num]
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))
            
            page_text = pytesseract.image_to_string(img)
            if page_text:
                text += page_text + "\n"
        doc.close()
        return text
    except Exception as e:
        logger.error(f"pytesseract OCR execution failed: {str(e)}")
        raise RuntimeError(f"OCR parsing failed. Error: {str(e)}")

def extract_text_from_pdf_pymupdf(file_bytes):
    if not PYMUPDF_AVAILABLE:
        raise ImportError("PyMuPDF is not available.")
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_to_parse = min(MAX_PAGES_TO_PARSE, len(doc))
    text = ""
    for i in range(pages_to_parse):
        text += doc[i].get_text()
    doc.close()
    return text

def extract_text_from_pdf_pdfplumber(file_bytes):
    if not PDFPLUMBER_AVAILABLE:
        raise ImportError("pdfplumber is not available.")
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages_to_parse = min(MAX_PAGES_TO_PARSE, len(pdf.pages))
        for i in range(pages_to_parse):
            page_text = pdf.pages[i].extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_pdf_pypdf(file_bytes):
    if not PYPDF_AVAILABLE:
        raise ImportError("pypdf is not available.")
    text = ""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    pages_to_parse = min(MAX_PAGES_TO_PARSE, len(reader.pages))
    for i in range(pages_to_parse):
        page_text = reader.pages[i].extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text(file_bytes, file_name):
    """
    Ingests PDF/DOCX files with robust hierarchical fallbacks and strict 3-page limits.
    """
    ext = os.path.splitext(file_name)[1].lower()
    
    if ext == '.pdf':
        # 1. Primary: PyMuPDF
        try:
            logger.info(f"Ingesting {file_name} using PyMuPDF (Capped at {MAX_PAGES_TO_PARSE} pages)...")
            text = extract_text_from_pdf_pymupdf(file_bytes)
            if text and text.strip():
                logger.info("Successfully extracted text using PyMuPDF.")
                return text, "PyMuPDF"
            else:
                logger.warning("PyMuPDF returned empty text. Trying pdfplumber...")
        except Exception as e:
            logger.error(f"PyMuPDF failed: {str(e)}. Trying pdfplumber fallback...")
            
        # 2. Fallback 1: pdfplumber
        try:
            logger.info("Ingesting using pdfplumber...")
            text = extract_text_from_pdf_pdfplumber(file_bytes)
            if text and text.strip():
                logger.info("Successfully extracted text using pdfplumber.")
                return text, "pdfplumber"
            else:
                logger.warning("pdfplumber returned empty text. Trying pypdf...")
        except Exception as e:
            logger.error(f"pdfplumber failed: {str(e)}. Trying pypdf fallback...")
            
        # 3. Fallback 2: pypdf
        try:
            logger.info("Ingesting using pypdf...")
            text = extract_text_from_pdf_pypdf(file_bytes)
            if text and text.strip():
                logger.info("Successfully extracted text using pypdf.")
                return text, "pypdf"
            else:
                logger.warning("pypdf returned empty text. Trying OCR fallback...")
        except Exception as e:
            logger.error(f"pypdf failed: {str(e)}. Trying OCR fallback...")
            
        # 4. Fallback 3: Tesseract OCR (Scanned PDF)
        try:
            text = extract_text_via_ocr(file_bytes)
            if text and text.strip():
                logger.info("Successfully extracted text via Tesseract OCR.")
                return text, "Tesseract-OCR"
            else:
                raise ValueError("OCR returned empty text.")
        except Exception as e:
            logger.error(f"OCR fallback failed: {str(e)}")
            raise RuntimeError(f"All extraction layers failed for PDF. File may be empty or corrupt. Last error: {str(e)}")
            
    elif ext in ['.docx', '.doc']:
        logger.info(f"Ingesting {file_name} using python-docx...")
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_list = []
            for paragraph in doc.paragraphs:
                text_list.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_list.append(cell.text)
            full_text = "\n".join(text_list)
            logger.info("Successfully extracted text using python-docx.")
            return full_text, "python-docx"
        except Exception as e:
            logger.error(f"python-docx failed: {str(e)}")
            raise RuntimeError(f"Failed to parse docx document. Error: {str(e)}")
            
    else:
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
            return text, "plain-text"
        except Exception as e:
            raise ValueError(f"Unsupported format: {ext} and plain-text fallback failed: {str(e)}")
